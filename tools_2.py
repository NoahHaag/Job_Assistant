import asyncio
import base64
import json
import os
import re
import uuid
from datetime import datetime, date

import docx
import requests
from PyPDF2 import PdfReader
from chromadb import Settings
import qrcode
from google.adk.tools import FunctionTool, google_search
from langchain_community.vectorstores import Chroma
from langchain_google_community import GmailToolkit
from langchain_google_community.gmail.utils import get_google_credentials, build_gmail_service

from get_embedding_function import get_embedding_function

from dotenv import load_dotenv

load_dotenv()


DOCUMENT_FOLDER = "documents"
SCRATCHPAD_FILE = "agent_scratchpad.txt"
JOB_APPLICATIONS_FILE = "job_applications.json"
COVER_LETTERS_FOLDER = "cover_letters"
JOB_OPPORTUNITIES_FILE = "job_opportunities.json"
SERPAPI_USAGE_FILE = "serpapi_usage.json"


# -----------------------------
# Write: Append text to file
# -----------------------------
def write_to_scratchpad(text: str):
    """
    Append text to the scratchpad with a timestamp.

    Args:
        text (str): The content to append.

    Returns:
        str: Confirmation message.
    """
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(SCRATCHPAD_FILE, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {text}\n\n")
        f.flush()  # ensure text is written immediately
    return "Scratchpad updated."


def read_scratchpad():
    """
    Read and return the full contents of the scratchpad.

    This function retrieves all text stored in the scratchpad file.
    If the file does not exist or is empty, it returns a placeholder
    message indicating the scratchpad is empty. It can be used by an
    agent to review prior notes or reasoning steps.

    Returns:
        str: The complete contents of the scratchpad, or a message
             indicating it is empty.

    Example:
        >>> read_scratchpad()
        'Consider checking the latest job openings.\n\nPrepare CV updates.\n\n'
    """
    if not os.path.exists(SCRATCHPAD_FILE):
        return "(scratchpad is empty)"
    with open(SCRATCHPAD_FILE, "r", encoding="utf-8") as f:
        return f.read()


# -----------------------------
# Wrapping in ADK FunctionTools
# -----------------------------
write_scratchpad_tool = FunctionTool(
    func=write_to_scratchpad
)

read_scratchpad_tool = FunctionTool(
    func=read_scratchpad
)


# ---------------------------------------------------------------------
# Job Application Tracker - JSON-based storage
# ---------------------------------------------------------------------

def _load_job_applications():
    """
    Load job applications from JSON file. Create file if it doesn't exist.
    
    Returns:
        dict: Job applications data structure
    """
    if not os.path.exists(JOB_APPLICATIONS_FILE):
        initial_data = {"applications": []}
        with open(JOB_APPLICATIONS_FILE, "w", encoding="utf-8") as f:
            json.dump(initial_data, f, indent=2)
        return initial_data
    
    try:
        with open(JOB_APPLICATIONS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError:
        print("[WARNING] Corrupted job_applications.json, creating backup and initializing fresh.")
        if os.path.exists(JOB_APPLICATIONS_FILE):
            backup_name = f"job_applications_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            os.rename(JOB_APPLICATIONS_FILE, backup_name)
        return {"applications": []}


def _save_job_applications(data):
    """
    Save job applications to JSON file.
    
    Args:
        data (dict): Job applications data structure
    """
    with open(JOB_APPLICATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def add_job_application(
    company: str,
    position: str,
    status: str = "applied",
    date_applied: str = None,
    application_deadline: str = None,
    job_description: str = "",
    next_action: str = "",
    notes: str = "",
    contacts: str = ""
):
    """
    Add a new job application to the tracker.
    
    Args:
        company (str): Company name
        position (str): Job title/position
        status (str): Application status - one of: applied, interview_scheduled, 
                     interviewed, rejected, offer, accepted. Defaults to "applied"
        date_applied (str, optional): Date applied in YYYY-MM-DD format. Defaults to today.
        application_deadline (str, optional): Application deadline in YYYY-MM-DD format
        job_description (str, optional): Full job description text
        next_action (str, optional): Next steps or follow-up actions
        notes (str, optional): Additional notes or context
        contacts (str, optional): Comma-separated list of contact emails
    
    Returns:
        str: Confirmation message with application ID
    
    Example:
        >>> add_job_application(
                company="Google",
                position="ML Engineer",
                status="applied",
                date_applied="2025-11-24",
                notes="Referred by John Doe"
            )
        "Job application added successfully! ID: abc-123-def"
    """
    valid_statuses = ["applied", "interview_scheduled", "interviewed", "rejected", "offer", "accepted"]
    if status not in valid_statuses:
        return f"Error: Invalid status '{status}'. Must be one of: {', '.join(valid_statuses)}"
    
    if date_applied is None:
        date_applied = date.today().isoformat()
    
    data = _load_job_applications()
    
    # Generate unique ID
    app_id = str(uuid.uuid4())[:8]
    
    # Process contacts
    contact_list = [c.strip() for c in contacts.split(",")] if contacts else []
    
    application = {
        "id": app_id,
        "company": company,
        "position": position,
        "status": status,
        "date_applied": date_applied,
        "application_deadline": application_deadline or "",
        "job_description": job_description,
        "cover_letter_generated": False,
        "next_action": next_action,
        "notes": notes,
        "contacts": contact_list,
        "last_updated": datetime.now().isoformat()
    }
    
    data["applications"].append(application)
    _save_job_applications(data)
    
    return f"✅ Job application added successfully!\n\nID: {app_id}\nCompany: {company}\nPosition: {position}\nStatus: {status}\nDate Applied: {date_applied}"


def update_job_application(
    application_id: str = None,
    company: str = None,
    status: str = None,
    next_action: str = None,
    notes: str = None,
    cover_letter_generated: bool = None
):
    """
    Update an existing job application.
    
    Args:
        application_id (str, optional): Application ID to update
        company (str, optional): Company name (used if application_id not provided)
        status (str, optional): New status
        next_action (str, optional): Updated next action
        notes (str, optional): Additional notes to append
        cover_letter_generated (bool, optional): Whether cover letter was generated
    
    Returns:
        str: Confirmation message with updated details
    
    Example:
        >>> update_job_application(company="Google", status="interview_scheduled")
        "Updated Google - ML Engineer to status: interview_scheduled"
    """
    if not application_id and not company:
        return "Error: Must provide either application_id or company name"
    
    data = _load_job_applications()
    
    # Find application
    app_to_update = None
    if application_id:
        for app in data["applications"]:
            if app["id"] == application_id:
                app_to_update = app
                break
    elif company:
        # Find most recent application for this company
        matching_apps = [app for app in data["applications"] if app["company"].lower() == company.lower()]
        if matching_apps:
            app_to_update = max(matching_apps, key=lambda x: x["last_updated"])
    
    if not app_to_update:
        return f"Error: No application found for {'ID ' + application_id if application_id else 'company ' + company}"
    
    # Update fields
    if status:
        valid_statuses = ["applied", "interview_scheduled", "interviewed", "rejected", "offer", "accepted"]
        if status not in valid_statuses:
            return f"Error: Invalid status '{status}'. Must be one of: {', '.join(valid_statuses)}"
        app_to_update["status"] = status
    
    if next_action is not None:
        app_to_update["next_action"] = next_action
    
    if notes is not None:
        # Append to existing notes
        if app_to_update["notes"]:
            app_to_update["notes"] += f"\n[{datetime.now().strftime('%Y-%m-%d')}] {notes}"
        else:
            app_to_update["notes"] = notes
    
    if cover_letter_generated is not None:
        app_to_update["cover_letter_generated"] = cover_letter_generated
    
    app_to_update["last_updated"] = datetime.now().isoformat()
    
    _save_job_applications(data)
    
    return f"✅ Updated: {app_to_update['company']} - {app_to_update['position']}\nStatus: {app_to_update['status']}\nLast Updated: {app_to_update['last_updated']}"


def get_job_applications(
    status: str = None,
    company: str = None,
    days_back: int = None,
    sort_by: str = "date_applied"
):
    """
    Query job applications with optional filters.
    
    Args:
        status (str, optional): Filter by status (e.g., "interview_scheduled")
        company (str, optional): Filter by company name (partial match)
        days_back (int, optional): Only show applications from last N days
        sort_by (str): Sort by field - one of: date_applied, last_updated, company, status
    
    Returns:
        str: Formatted list of matching applications
    
    Example:
        >>> get_job_applications(status="interview_scheduled")
        "Found 2 applications:
        1. Google - ML Engineer (interview_scheduled) - Applied: 2025-11-20
        2. Meta - Research Scientist (interview_scheduled) - Applied: 2025-11-22"
    """
    data = _load_job_applications()
    applications = data["applications"]
    
    # Apply filters
    if status:
        applications = [app for app in applications if app["status"] == status]
    
    if company:
        applications = [app for app in applications if company.lower() in app["company"].lower()]
    
    if days_back:
        cutoff_date = (date.today() - __import__('datetime').timedelta(days=days_back)).isoformat()
        applications = [app for app in applications if app["date_applied"] >= cutoff_date]
    
    if not applications:
        filter_desc = []
        if status:
            filter_desc.append(f"status={status}")
        if company:
            filter_desc.append(f"company={company}")
        if days_back:
            filter_desc.append(f"last {days_back} days")
        
        filter_str = " with filters: " + ", ".join(filter_desc) if filter_desc else ""
        return f"No applications found{filter_str}."
    
    # Sort
    sort_fields = {
        "date_applied": lambda x: x["date_applied"],
        "last_updated": lambda x: x["last_updated"],
        "company": lambda x: x["company"],
        "status": lambda x: x["status"]
    }
    
    if sort_by in sort_fields:
        applications = sorted(applications, key=sort_fields[sort_by], reverse=True)
    
    # Format output
    result = [f"📊 Found {len(applications)} application(s):\n"]
    
    for i, app in enumerate(applications, 1):
        deadline_str = f" | Deadline: {app['application_deadline']}" if app['application_deadline'] else ""
        next_action_str = f"\n   Next: {app['next_action']}" if app['next_action'] else ""
        cover_letter_str = " ✓ Cover Letter" if app.get('cover_letter_generated') else ""
        
        result.append(
            f"{i}. {app['company']} - {app['position']}{cover_letter_str}\n"
            f"   ID: {app['id']} | Status: {app['status']} | Applied: {app['date_applied']}{deadline_str}{next_action_str}"
        )
    
    return "\n\n".join(result)


def delete_job_application(application_id: str = None, company: str = None):
    """
    Delete a job application from the tracker.
    
    Args:
        application_id (str, optional): Application ID to delete
        company (str, optional): Company name (deletes most recent if multiple exist)
    
    Returns:
        str: Confirmation message
    
    Example:
        >>> delete_job_application(company="OldCompany")
        "Deleted application: OldCompany - Engineer"
    """
    if not application_id and not company:
        return "Error: Must provide either application_id or company name"
    
    data = _load_job_applications()
    
    # Find and remove application
    original_count = len(data["applications"])
    
    if application_id:
        data["applications"] = [app for app in data["applications"] if app["id"] != application_id]
        identifier = f"ID {application_id}"
    elif company:
        # Find most recent to delete
        matching_apps = [app for app in data["applications"] if app["company"].lower() == company.lower()]
        if matching_apps:
            app_to_delete = max(matching_apps, key=lambda x: x["last_updated"])
            data["applications"] = [app for app in data["applications"] if app["id"] != app_to_delete["id"]]
            identifier = f"{app_to_delete['company']} - {app_to_delete['position']}"
        else:
            return f"Error: No application found for company '{company}'"
    
    if len(data["applications"]) == original_count:
        return f"Error: Application not found"
    
    _save_job_applications(data)
    return f"✅ Deleted application: {identifier}"


# ---------------------------------------------------------------------
# Cold Email Tracker - Track outreach to professors/researchers
# ---------------------------------------------------------------------

COLD_EMAILS_FILE = "cold_emails.json"

def _load_cold_emails():
    """Load cold emails from JSON file. Create file if it doesn't exist."""
    if not os.path.exists(COLD_EMAILS_FILE):
        initial_data = {"emails": []}
        with open(COLD_EMAILS_FILE, "w", encoding="utf-8") as f:
            json.dump(initial_data, f, indent=2)
        return initial_data
    
    try:
        with open(COLD_EMAILS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError:
        print("[WARNING] Corrupted cold_emails.json, creating backup.")
        if os.path.exists(COLD_EMAILS_FILE):
            backup_name = f"cold_emails_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            os.rename(COLD_EMAILS_FILE, backup_name)
        return {"emails": []}


def _save_cold_emails(data):
    """Save cold emails to JSON file."""
    with open(COLD_EMAILS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def add_cold_email(
    recipient_name: str,
    recipient_email: str,
    institution: str = "",
    subject: str = "",
    purpose: str = "",
    date_sent: str = None,
    notes: str = "",
    referred_by: str = "",
    connection_strength: int = 1
):
    """
    Add a new cold email to the tracker, or update existing one if recipient_email already exists.
    
    Args:
        recipient_name (str): Name of the person you emailed
        recipient_email (str): Their email address
        institution (str, optional): University or institution
        subject (str, optional): Email subject line
        purpose (str, optional): Purpose (e.g., "PhD opportunity")
        date_sent (str, optional): Date sent (YYYY-MM-DD). Defaults to today.
        notes (str, optional): Additional notes
        referred_by (str, optional): Name or ID of person who referred you
        connection_strength (int, optional): 1-5 rating of relationship strength (default: 1)
    
    Returns:
        str: Confirmation message with email ID
    """
    if date_sent is None:
        date_sent = date.today().isoformat()
    
    data = _load_cold_emails()
    
    # Check if an email to this recipient already exists
    existing_email = None
    for email in data["emails"]:
        if email["recipient_email"].lower() == recipient_email.lower():
            existing_email = email
            break
    
    if existing_email:
        # Update existing entry
        updated_fields = []
        
        # Update recipient_name if provided and different
        if recipient_name and recipient_name != existing_email["recipient_name"]:
            existing_email["recipient_name"] = recipient_name
            updated_fields.append("name")
        
        # Update institution if provided and different
        if institution and institution != existing_email["institution"]:
            existing_email["institution"] = institution
            updated_fields.append("institution")
        
        # Update subject if provided
        if subject and subject != existing_email["subject"]:
            existing_email["subject"] = subject
            updated_fields.append("subject")
        
        # Update purpose if provided
        if purpose and purpose != existing_email["purpose"]:
            existing_email["purpose"] = purpose
            updated_fields.append("purpose")
        
        # Add to follow-up dates if this is a new send
        if date_sent not in existing_email.get("follow_up_dates", []):
            existing_email["follow_up_dates"].append(date_sent)
            updated_fields.append("follow-up date")
        
        # Append notes if provided
        if notes:
            if existing_email["notes"]:
                existing_email["notes"] += f"\n[{datetime.now().strftime('%Y-%m-%d')}] {notes}"
            else:
                existing_email["notes"] = notes
            updated_fields.append("notes")
        
        # Update referred_by if provided
        if referred_by and referred_by != existing_email.get("referred_by", ""):
            existing_email["referred_by"] = referred_by
            updated_fields.append("referred_by")

        # Update connection_strength if provided
        if connection_strength != existing_email.get("connection_strength", 1):
            existing_email["connection_strength"] = connection_strength
            updated_fields.append("connection_strength")
        
        existing_email["last_updated"] = datetime.now().isoformat()
        _save_cold_emails(data)
        
        fields_str = ", ".join(updated_fields) if updated_fields else "no new information"
        return f"🔄 Updated existing email!\n\nID: {existing_email['id']}\nRecipient: {existing_email['recipient_name']} ({recipient_email})\nInstitution: {existing_email['institution'] or 'N/A'}\nUpdated: {fields_str}"
    
    else:
        # Create new entry
        email_id = str(uuid.uuid4())[:8]
        
        cold_email = {
            "id": email_id,
            "recipient_name": recipient_name,
            "recipient_email": recipient_email,
            "institution": institution,
            "subject": subject,
            "purpose": purpose,
            "date_sent": date_sent,
            "status": "sent",
            "response_date": None,
            "follow_up_dates": [],
            "notes": notes,
            "last_updated": datetime.now().isoformat(),
            "referred_by": referred_by,
            "connection_strength": connection_strength
        }
        
        data["emails"].append(cold_email)
        _save_cold_emails(data)
        
        return f"✅ Cold email tracked!\n\nID: {email_id}\nRecipient: {recipient_name} ({recipient_email})\nInstitution: {institution or 'N/A'}\nDate: {date_sent}"


def update_cold_email(
    email_id: str = None,
    recipient_email: str = None,
    recipient_name: str = None,
    status: str = None,
    response_date: str = None,
    follow_up_sent: bool = False,
    notes: str = None,
    referred_by: str = None,
    connection_strength: int = None
):
    """
    Update an existing cold email record.
    
    Args:
        email_id (str, optional): Email ID to update
        recipient_email (str, optional): Recipient's email (if no email_id)
        recipient_name (str, optional): Recipient's name (if no email_id/email) - partial match
        status (str, optional): New status (sent, responded, no_response, follow_up_sent)
        response_date (str, optional): Date they responded (YYYY-MM-DD)
        follow_up_sent (bool): If True, adds today to follow_up_dates
        follow_up_sent (bool): If True, adds today to follow_up_dates
        notes (str, optional): Additional notes to append
        referred_by (str, optional): Update referrer
        connection_strength (int, optional): Update connection strength (1-5)
    
    Returns:
        str: Confirmation message
    """
    print(f"[DEBUG] update_cold_email called with: name={recipient_name}, email={recipient_email}, id={email_id}")
    if not email_id and not recipient_email and not recipient_name:
        return "Error: Must provide email_id, recipient_email, or recipient_name"
    
    valid_statuses = ["sent", "responded", "no_response", "follow_up_sent"]
    if status and status not in valid_statuses:
        return f"Error: Invalid status. Must be one of: {', '.join(valid_statuses)}"
    
    data = _load_cold_emails()
    
    # Find email
    email_to_update = None
    if email_id:
        for email in data["emails"]:
            if email["id"] == email_id:
                email_to_update = email
                break
    elif recipient_email:
        matching_emails = [e for e in data["emails"] if e["recipient_email"].lower() == recipient_email.lower()]
        if matching_emails:
            email_to_update = max(matching_emails, key=lambda x: x["last_updated"])
    elif recipient_name:
        # Partial match on name
        matching_emails = [e for e in data["emails"] if recipient_name.lower() in e["recipient_name"].lower()]
        if len(matching_emails) == 0:
            return f"Error: No email found matching name '{recipient_name}'"
        elif len(matching_emails) > 1:
            names = [f"{e['recipient_name']} ({e['recipient_email']})" for e in matching_emails]
            return f"Error: Multiple emails found matching '{recipient_name}': {', '.join(names)}. Please be more specific."
        else:
            email_to_update = matching_emails[0]
    
    if not email_to_update:
        return "Error: No email found"
    
    # Update fields
    if status:
        email_to_update["status"] = status
    if response_date:
        email_to_update["response_date"] = response_date
        if not status:
            email_to_update["status"] = "responded"
    if follow_up_sent:
        email_to_update["follow_up_dates"].append(date.today().isoformat())
        if not status:
            email_to_update["status"] = "follow_up_sent"
    if notes:
        if email_to_update["notes"]:
            email_to_update["notes"] += f"\n[{datetime.now().strftime('%Y-%m-%d')}] {notes}"
        else:
            email_to_update["notes"] = notes
    
    if referred_by is not None:
        email_to_update["referred_by"] = referred_by
    
    if connection_strength is not None:
        email_to_update["connection_strength"] = connection_strength
    
    email_to_update["last_updated"] = datetime.now().isoformat()
    _save_cold_emails(data)
    
    return f"✅ Updated: {email_to_update['recipient_name']} ({email_to_update['recipient_email']})\nStatus: {email_to_update['status']}"


def query_cold_emails(
    status: str = None,
    institution: str = None,
    recipient_name: str = None,
    days_back: int = None,
    awaiting_response: bool = False
):
    """
    Query cold emails with optional filters.
    
    Args:
        status (str, optional): Filter by status
        institution (str, optional): Filter by institution (partial match)
        recipient_name (str, optional): Filter by name (partial match)
        days_back (int, optional): Show emails from last N days
        awaiting_response (bool): Show only sent emails with no response
    
    Returns:
        str: Formatted list of matching emails
    """
    data = _load_cold_emails()
    emails = data["emails"]
    
    # Apply filters
    if status:
        emails = [e for e in emails if e["status"] == status]
    if institution:
        emails = [e for e in emails if institution.lower() in e["institution"].lower()]
    if recipient_name:
        emails = [e for e in emails if recipient_name.lower() in e["recipient_name"].lower()]
    if days_back:
        cutoff_date = (date.today() - __import__('datetime').timedelta(days=days_back)).isoformat()
        emails = [e for e in emails if e["date_sent"] >= cutoff_date]
    if awaiting_response:
        emails = [e for e in emails if e["status"] in ["sent", "follow_up_sent"] and not e["response_date"]]
    
    if not emails:
        return "No cold emails found with those filters."
    
    # Sort by date
    emails = sorted(emails, key=lambda x: x["date_sent"], reverse=True)
    
    # Format output
    result = [f"📧 Found {len(emails)} cold email(s):\n"]
    
    for i, email in enumerate(emails, 1):
        inst = f" ({email['institution']})" if email['institution'] else ""
        purp = f" - {email['purpose']}" if email['purpose'] else ""
        resp = f"\n   ✅ Responded: {email['response_date']}" if email['response_date'] else ""
        follows = f"\n   🔄 Follow-ups: {len(email['follow_up_dates'])}" if email['follow_up_dates'] else ""
        
        result.append(
            f"{i}. {email['recipient_name']}{inst}{purp}\n"
            f"   {email['recipient_email']} | {email['status']} | Sent: {email['date_sent']}{resp}{follows}"
        )
    
    return "\n\n".join(result)


def generate_network_graph():
    """
    Generate a Mermaid.js graph visualization of your professional network.
    
    This tool reads the cold email tracker and visualizes relationships between:
    - You (Central Node)
    - Contacts (People you've emailed)
    - Institutions (Where they work)
    - Referrals (Who introduced whom)
    
    Returns:
        str: Mermaid.js graph syntax to be rendered
    """
    data = _load_cold_emails()
    emails = data["emails"]
    
    # Start Mermaid graph
    graph = ["graph TD"]
    
    # Style definitions
    graph.append("    %% Styles")
    graph.append("    classDef person fill:#e1f5fe,stroke:#01579b,stroke-width:2px,rx:10,ry:10;")
    graph.append("    classDef institution fill:#f3e5f5,stroke:#4a148c,stroke-width:2px,rx:5,ry:5;")
    graph.append("    classDef responded fill:#c8e6c9,stroke:#2e7d32,stroke-width:3px;")
    graph.append("    classDef no_response fill:#ffcdd2,stroke:#c62828,stroke-width:1px;")
    
    # Nodes and Edges
    institutions = set()
    
    # Central Node (User)
    graph.append("    User(You) --> Network")
    graph.append("    style User fill:#fff9c4,stroke:#fbc02d,stroke-width:4px")
    
    for email in emails:
        # Sanitize IDs for Mermaid (remove special chars)
        p_id = f"P_{email['id']}"
        p_name = email['recipient_name'].replace("(", "").replace(")", "").replace("[", "").replace("]", "")
        p_inst = email['institution'].strip() if email['institution'] else "Unknown"
        i_id = f"I_{abs(hash(p_inst))}"
        
        # Determine style based on status
        style_class = "person"
        if email.get("response_date"):
            style_class += ",responded"
        else:
            style_class += ",no_response"
            
        # Add Person Node
        graph.append(f"    {p_id}({p_name}):::{style_class}")
        
        # Add Institution Node and Edge
        if p_inst:
            if p_inst not in institutions:
                graph.append(f"    {i_id}[{p_inst}]:::institution")
                institutions.add(p_inst)
            graph.append(f"    {p_id} -- Works at --> {i_id}")
            
        # Add Referral Edge
        referrer = email.get("referred_by")
        if referrer:
            # Try to find referrer in existing contacts
            referrer_id = None
            for e in emails:
                if e['recipient_name'].lower() == referrer.lower() or e['id'] == referrer:
                    referrer_id = f"P_{e['id']}"
                    break
            
            if referrer_id:
                graph.append(f"    {referrer_id} -- Referred --> {p_id}")
            else:
                # External referrer
                r_id = f"R_{abs(hash(referrer))}"
                graph.append(f"    {r_id}({referrer})")
                graph.append(f"    {r_id} -- Referred --> {p_id}")

    return "\n".join(graph)


# ---------------------------------------------------------------------
# Job Opportunities Scraper - SerpAPI with usage tracking
# ---------------------------------------------------------------------

def _load_job_opportunities():
    """Load job opportunities from JSON file. Create file if it doesn't exist."""
    if not os.path.exists(JOB_OPPORTUNITIES_FILE):
        initial_data = {"opportunities": []}
        with open(JOB_OPPORTUNITIES_FILE, "w", encoding="utf-8") as f:
            json.dump(initial_data, f, indent=2)
        return initial_data
    
    try:
        with open(JOB_OPPORTUNITIES_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError:
        print("[WARNING] Corrupted job_opportunities.json, creating backup.")
        if os.path.exists(JOB_OPPORTUNITIES_FILE):
            backup_name = f"job_opportunities_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            os.rename(JOB_OPPORTUNITIES_FILE, backup_name)
        return {"opportunities": []}


def _save_job_opportunities(data):
    """Save job opportunities to JSON file."""
    with open(JOB_OPPORTUNITIES_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------
# Job Fair Tools
# ---------------------------------------------------------------------

def elevator_pitch_tool(company_name: str, job_description: str = ""):
    """
    Generate a 30-second elevator pitch tailored to a specific company.
    
    Args:
        company_name (str): The company you are speaking to.
        job_description (str, optional): Key details about the role.
    
    Returns:
        str: A tailored elevator pitch.
    """
    # In a real scenario, we would read the CV here. For now, we'll assume the agent passes relevant context 
    # or we rely on the LLM's knowledge of the user from the session.
    # But since this is a tool, we should probably read the CV to be safe/robust if called standalone.
    
    # Read CV and Resume
    cv_text = read_document("Professional Curriculum Vitae.docx")
    resume_text = read_document("Resume.docx")
    
    context = ""
    if not cv_text.startswith("Error"):
        context += f"\nCV Content:\n{cv_text}\n"
    if not resume_text.startswith("Error"):
        context += f"\nResume Content:\n{resume_text}\n"
        
    return f"Please generate an elevator pitch for {company_name}. Key JD points: {job_description}. \n\nMy Background Context:\n{context}"


def company_brief_tool(company_name: str):
    """
    Generate a brief summary of a company using Google Search.
    
    Args:
        company_name (str): Name of the company.
        
    Returns:
        str: Summary of the company.
    """
    # We can use the existing google_search tool function if available, 
    # or just return a prompt for the agent to do it.
    # Given the GHA requirement, we should probably implement the search here if we want it to be "one click".
    # But `google_search` is an ADK tool.
    
    return f"Please search for recent news, mission, and values for {company_name} and create a 1-page brief with 3 strategic questions."


def qr_code_tool(data: str, filename: str = "qrcode.png"):
    """
    Generate a QR code and save it to a file.
    
    Args:
        data (str): The text or URL to encode.
        filename (str): Output filename (default: qrcode.png).
        
    Returns:
        str: Path to the generated QR code.
    """
    # Ensure filename is in a safe place, e.g., artifacts or root
    # For GHA, root is fine.
    
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(data)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")
    
    # Save to current directory or specific folder
    output_path = os.path.abspath(filename)
    img.save(output_path)
    
    return f"QR Code generated at: {output_path}"


def portfolio_export_tool():
    """
    Export job search stats and network graph to a Markdown summary.
    
    Returns:
        str: Markdown content.
    """
    # 1. Get Job Stats
    apps = _load_job_applications()["applications"]
    total_apps = len(apps)
    interviews = len([a for a in apps if "interview" in a["status"]])
    offers = len([a for a in apps if "offer" in a["status"]])
    
    # 2. Get Network Stats
    emails = _load_cold_emails()["emails"]
    total_emails = len(emails)
    responses = len([e for e in emails if e["response_date"]])
    
    # 3. Generate Markdown
    md = f"""# 🚀 Job Search Portfolio

## 📊 At a Glance
- **Applications**: {total_apps}
- **Interviews**: {interviews}
- **Offers**: {offers}
- **Network Outreach**: {total_emails} emails sent ({responses} responses)

## 🕸️ Network Graph
(Rendered in Mermaid.js below)

```mermaid
{generate_network_graph()}
```

## 🛠️ Built With
This portfolio is managed by my **AI Agent** powered by Gemini.
"""
    return md



def _load_serpapi_usage():
    """Load SerpAPI usage tracker."""
    if not os.path.exists(SERPAPI_USAGE_FILE):
        initial_data = {"monthly_limit": 200, "searches": []}
        with open(SERPAPI_USAGE_FILE, "w", encoding="utf-8") as f:
            json.dump(initial_data, f, indent=2)
        return initial_data
    
    try:
        with open(SERPAPI_USAGE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError:
        print("[WARNING] Corrupted serpapi_usage.json, resetting.")
        return {"monthly_limit": 200, "searches": []}


def _save_serpapi_usage(data):
    """Save SerpAPI usage tracker."""
    with open(SERPAPI_USAGE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


def _count_searches_this_month(usage_data):
    """Count searches in current month."""
    current_month = datetime.now().strftime("%Y-%m")
    return sum(1 for s in usage_data["searches"] if s["date"].startswith(current_month))


def _log_serpapi_search(query, results_count):
    """Log a SerpAPI search."""
    usage_data = _load_serpapi_usage()
    usage_data["searches"].append({
        "date": datetime.now().isoformat(),
        "query": query,
        "results": results_count
    })
    _save_serpapi_usage(usage_data)


def _is_duplicate_job(job, existing_jobs):
    """Check if a job is a duplicate based on company, title, and location."""
    for existing in existing_jobs:
        if (existing.get("company", "").lower() == job.get("company", "").lower() and
            existing.get("title", "").lower() == job.get("title", "").lower() and
            existing.get("location", "").lower() == job.get("location", "").lower()):
            return True
    return False


def search_jobs_serpapi(
    query: str,
    location: str = "",
    date_posted: str = "today",
    max_results: int = 10
):
    """
    Core SerpAPI job search function.
    
    Args:
        query: Job search query (e.g., "Marine Scientist")
        location: Location (e.g., "Florida, USA")
        date_posted: Filter by date - "today", "3days", "week", "month"
        max_results: Number of results to return (max 10 for free tier)
    
    Returns:
        List of job dictionaries
    """
    api_key = os.getenv("SERPAPI_KEY")
    
    if not api_key:
        return {"error": "SERPAPI_KEY not found in environment variables. Please add it to your .env file."}
    
    url = "https://serpapi.com/search"
    
    # Map date_posted to SerpAPI chips format
    date_mapping = {
        "today": "date_posted:today",
        "3days": "date_posted:3days",
        "week": "date_posted:week",
        "month": "date_posted:month"
    }
    chips = date_mapping.get(date_posted, "date_posted:week")
    
    params = {
        "engine": "google_jobs",
        "q": query,
        "location": location,
        "api_key": api_key,
        "chips": chips,
        "num": min(max_results, 10)  # Free tier limit
    }
    
    try:
        response = requests.get(url, params=params, timeout=10)
        response.raise_for_status()
        results = response.json()
        
        jobs = []
        for job in results.get("jobs_results", []):
            jobs.append({
                "title": job.get("title", ""),
                "company": job.get("company_name", ""),
                "location": job.get("location", ""),
                "link": job.get("share_link", ""),
                "description": job.get("description", "")[:500],  # Truncate long descriptions
                "via": job.get("via", ""),
                "date_posted": job.get("detected_extensions", {}).get("posted_at", ""),
                "salary": job.get("detected_extensions", {}).get("salary", "")
            })
        
        return jobs
    
    except requests.exceptions.HTTPError as e:
        error_msg = f"SerpAPI request failed: {str(e)}"
        if e.response is not None:
            try:
                error_msg += f"\nResponse: {e.response.text}"
            except:
                error_msg += "\nResponse: (could not decode response text)"
        return {"error": error_msg}
    except requests.exceptions.RequestException as e:
        return {"error": f"SerpAPI request failed: {str(e)}"}
    except Exception as e:
        return {"error": f"Unexpected error: {str(e)}"}


def search_jobs(
    query: str,
    location: str = "",
    date_posted: str = "week",
    max_results: int = 10,
    usage_limit: int = 95,
    save_results: bool = True
):
    """
    Search for jobs using SerpAPI with built-in usage tracking and storage.
    
    Args:
        query: Job search query (e.g., "Marine Scientist", "Research Biologist")
        location: Location (e.g., "Florida, USA", "Remote")
        date_posted: Filter by date - "today", "3days", "week", "month" (default: "week")
        max_results: Number of results to return (max 10, default: 10)
        usage_limit: Stop searching when this many searches reached (default: 95)
        save_results: Save jobs to job_opportunities.json (default: True)
    
    Returns:
        dict with keys: "jobs", "new_jobs_count", "usage_stats", "warning"
    
    Example:
        >>> search_jobs("Marine Scientist", "Florida", "week")
        {
            "jobs": [...],
            "new_jobs_count": 5,
            "usage_stats": {"used": 23, "limit": 100, "remaining": 77},
            "warning": None
        }
    """
    # Check current usage
    usage_data = _load_serpapi_usage()
    searches_this_month = _count_searches_this_month(usage_data)
    remaining = usage_data["monthly_limit"] - searches_this_month
    
    # Check if we're over the limit
    if searches_this_month >= usage_limit:
        return {
            "jobs": [],
            "new_jobs_count": 0,
            "usage_stats": {
                "used": searches_this_month,
                "limit": usage_data["monthly_limit"],
                "remaining": remaining
            },
            "warning": f"⚠️ Usage limit reached ({searches_this_month}/{usage_limit}). Skipping search to preserve quota. Resets on the 1st of next month."
        }
    
    # Perform search
    jobs = search_jobs_serpapi(query, location, date_posted, max_results)
    
    # Check for errors
    if isinstance(jobs, dict) and "error" in jobs:
        return {
            "jobs": [],
            "new_jobs_count": 0,
            "usage_stats": {
                "used": searches_this_month,
                "limit": usage_data["monthly_limit"],
                "remaining": remaining
            },
            "warning": f"❌ Search failed: {jobs['error']}"
        }
    
    # Log the search
    _log_serpapi_search(query, len(jobs))
    searches_this_month += 1
    remaining -= 1
    
    # Save results if requested
    new_jobs_count = 0
    if save_results and jobs:
        opportunities_data = _load_job_opportunities()
        existing_jobs = opportunities_data["opportunities"]
        
        for job in jobs:
            # Skip duplicates
            if not _is_duplicate_job(job, existing_jobs):
                job_id = str(uuid.uuid4())[:8]
                opportunity = {
                    "id": job_id,
                    "title": job["title"],
                    "company": job["company"],
                    "location": job["location"],
                    "link": job["link"],
                    "description": job["description"],
                    "via": job["via"],
                    "date_posted": job["date_posted"],
                    "salary": job["salary"],
                    "date_discovered": datetime.now().isoformat(),
                    "search_query": f"{query} {location}".strip(),
                    "applied": False
                }
                opportunities_data["opportunities"].append(opportunity)
                new_jobs_count += 1
        
        if new_jobs_count > 0:
            _save_job_opportunities(opportunities_data)
    
    # Generate warning if approaching limit
    warning = None
    if searches_this_month >= usage_limit * 0.8:  # 80% threshold
        warning = f"⚠️ Approaching usage limit: {searches_this_month}/{usage_limit} searches used this month"
    
    return {
        "jobs": jobs,
        "new_jobs_count": new_jobs_count,
        "usage_stats": {
            "used": searches_this_month,
            "limit": usage_data["monthly_limit"],
            "remaining": remaining
        },
        "warning": warning,
        "message": f"✅ Found {len(jobs)} jobs, {new_jobs_count} new opportunities saved!" if new_jobs_count > 0 else f"Found {len(jobs)} jobs (all previously discovered)."
    }


def get_job_opportunities(
    days_back: int = None,
    company: str = None,
    title: str = None,
    sort_by: str = "date_discovered"
):
    """
    Query saved job opportunities with optional filters.
    
    Args:
        days_back: Only show jobs discovered in last N days
        company: Filter by company name (partial match)
        title: Filter by job title (partial match)
        sort_by: Sort by field - "date_discovered", "company", "title" (default: "date_discovered")
    
    Returns:
        Formatted string with matching job opportunities
    
    Example:
        >>> get_job_opportunities(days_back=7, title="Marine")
        "📋 Found 3 job opportunities:
        1. Marine Scientist - NOAA (Key West, FL)
           Discovered: Nov 27 | Via: LinkedIn
           Link: https://..."
    """
    data = _load_job_opportunities()
    opportunities = data["opportunities"]
    
    # Apply filters
    if days_back:
        cutoff_date = (datetime.now() - __import__('datetime').timedelta(days=days_back)).isoformat()
        opportunities = [opp for opp in opportunities if opp["date_discovered"] >= cutoff_date]
    
    if company:
        opportunities = [opp for opp in opportunities if company.lower() in opp["company"].lower()]
    
    if title:
        opportunities = [opp for opp in opportunities if title.lower() in opp["title"].lower()]
    
    if not opportunities:
        return "No job opportunities found with those filters."
    
    # Sort
    sort_fields = {
        "date_discovered": lambda x: x["date_discovered"],
        "company": lambda x: x["company"],
        "title": lambda x: x["title"]
    }
    
    if sort_by in sort_fields:
        opportunities = sorted(opportunities, key=sort_fields[sort_by], reverse=True)
    
    # Format output
    result = [f"📋 Found {len(opportunities)} job opportunity/opportunities:\n"]
    
    for i, opp in enumerate(opportunities, 1):
        date_str = datetime.fromisoformat(opp["date_discovered"]).strftime("%b %d")
        salary_str = f" | {opp['salary']}" if opp['salary'] else ""
        applied_str = " ✅ APPLIED" if opp.get('applied') else ""
        
        result.append(
            f"{i}. {opp['title']} - {opp['company']}{applied_str}\n"
            f"   📍 {opp['location']} | Discovered: {date_str} | Via: {opp['via']}{salary_str}\n"
            f"   🔗 {opp['link']}\n"
            f"   ID: {opp['id']}"
        )
    
    return "\n\n".join(result)


def get_serpapi_usage_report():
    """
    Get detailed SerpAPI usage report.
    
    Returns:
        Formatted string with usage stats and recent search history
    
    Example:
        >>> get_serpapi_usage_report()
        "📊 SerpAPI Usage Report
        Month: November 2025
        Used: 23/100 (23.0%)
        Remaining: 77 searches
        
        Recent Searches:
          - Nov 29, 12:30 PM: 'Marine Scientist Florida' (8 results)
          - Nov 29, 09:15 AM: 'Research Biologist Remote' (5 results)"
    """
    usage_data = _load_serpapi_usage()
    searches_this_month = _count_searches_this_month(usage_data)
    remaining = usage_data["monthly_limit"] - searches_this_month
    
    # Calculate percentage
    percentage = (searches_this_month / usage_data["monthly_limit"]) * 100
    
    report = [
        "📊 SerpAPI Usage Report",
        f"Month: {datetime.now().strftime('%B %Y')}",
        f"Used: {searches_this_month}/{usage_data['monthly_limit']} ({percentage:.1f}%)",
        f"Remaining: {remaining} searches",
        ""
    ]
    
    # Recent searches
    current_month = datetime.now().strftime("%Y-%m")
    recent = [s for s in usage_data["searches"] if s["date"].startswith(current_month)][-5:]
    
    if recent:
        report.append("Recent Searches:")
        for s in recent:
            date_str = datetime.fromisoformat(s["date"]).strftime("%b %d, %I:%M %p")
            report.append(f"  - {date_str}: '{s['query']}' ({s['results']} results)")
    else:
        report.append("No searches this month yet.")
    
    return "\n".join(report)


def delete_job_opportunity(job_id: str):
    """
    Delete a job opportunity from the tracker.
    
    Args:
        job_id: Job opportunity ID to delete
    
    Returns:
        Confirmation message
    
    Example:
        >>> delete_job_opportunity("abc12345")
        "✅ Deleted job opportunity: Marine Scientist - NOAA"
    """
    data = _load_job_opportunities()
    
    # Find and remove job
    original_count = len(data["opportunities"])
    job_to_delete = None
    
    for opp in data["opportunities"]:
        if opp["id"] == job_id:
            job_to_delete = opp
            break
    
    if not job_to_delete:
        return f"Error: Job opportunity with ID '{job_id}' not found"
    
    data["opportunities"] = [opp for opp in data["opportunities"] if opp["id"] != job_id]
    
    if len(data["opportunities"]) == original_count:
        return f"Error: Failed to delete job opportunity"
    
    _save_job_opportunities(data)
    return f"✅ Deleted job opportunity: {job_to_delete['title']} - {job_to_delete['company']}"


# Wrap as ADK FunctionTools
job_tracker_add_tool = FunctionTool(func=add_job_application)
job_tracker_update_tool = FunctionTool(func=update_job_application)
job_tracker_query_tool = FunctionTool(func=get_job_applications)
job_tracker_delete_tool = FunctionTool(func=delete_job_application)

cold_email_add_tool = FunctionTool(func=add_cold_email)
cold_email_update_tool = FunctionTool(func=update_cold_email)
cold_email_query_tool = FunctionTool(func=query_cold_emails)

network_graph_tool = FunctionTool(func=generate_network_graph)

# Job Search / Opportunities Tools
job_search_tool = FunctionTool(func=search_jobs)
job_opportunities_query_tool = FunctionTool(func=get_job_opportunities)
serpapi_usage_tool = FunctionTool(func=get_serpapi_usage_report)
job_opportunity_delete_tool = FunctionTool(func=delete_job_opportunity)


# ---------------------------------------------------------------------
# Google Scholar Search - SerpAPI with usage tracking
# ---------------------------------------------------------------------

def search_google_scholar_serpapi(
    query: str,
    as_ylo: int = None,
    as_yhi: int = None,
    max_results: int = 10
):
    """
    Core SerpAPI Google Scholar search function.
    
    Args:
        query: Search query (e.g., "coral bleaching")
        as_ylo: Start year (optional)
        as_yhi: End year (optional)
        max_results: Number of results to return (max 10 for free tier)
    
    Returns:
        List of result dictionaries
    """
    api_key = os.getenv("SERPAPI_KEY")
    
    if not api_key:
        return {"error": "SERPAPI_KEY not found in environment variables."}
    
    url = "https://serpapi.com/search"
    
    params = {
        "engine": "google_scholar",
        "q": query,
        "api_key": api_key,
        "num": min(max_results, 10)
    }
    
    if as_ylo:
        params["as_ylo"] = as_ylo
    if as_yhi:
        params["as_yhi"] = as_yhi
    
    try:
        response = requests.get(url, params=params, timeout=10)
        response.raise_for_status()
        results = response.json()
        
        organic_results = []
        for result in results.get("organic_results", []):
            organic_results.append({
                "title": result.get("title", ""),
                "link": result.get("link", ""),
                "snippet": result.get("snippet", ""),
                "publication_info": result.get("publication_info", {}).get("summary", ""),
                "cited_by": result.get("inline_links", {}).get("cited_by", {}).get("total", 0),
                "year": result.get("publication_info", {}).get("summary", "").split(" - ")[-1] if " - " in result.get("publication_info", {}).get("summary", "") else ""
            })
        
        return organic_results
    
    except requests.exceptions.RequestException as e:
        return {"error": f"SerpAPI request failed: {str(e)}"}
    except Exception as e:
        return {"error": f"Unexpected error: {str(e)}"}


def search_google_scholar(
    query: str,
    year_start: int = None,
    year_end: int = None,
    max_results: int = 5,
    usage_limit: int = 95
):
    """
    Search Google Scholar using SerpAPI with usage tracking.
    
    Args:
        query: Search query (e.g., "Acropora cervicornis restoration")
        year_start: Start year (optional)
        year_end: End year (optional)
        max_results: Number of results (default: 5)
        usage_limit: Stop searching when this many monthly searches reached (default: 95)
    
    Returns:
        Formatted string with search results
    """
    # Check usage
    usage_data = _load_serpapi_usage()
    searches_this_month = _count_searches_this_month(usage_data)
    
    if searches_this_month >= usage_limit:
        return f"⚠️ Usage limit reached ({searches_this_month}/{usage_data['monthly_limit']}). Search skipped."
    
    # Perform search
    results = search_google_scholar_serpapi(query, year_start, year_end, max_results)
    
    if isinstance(results, dict) and "error" in results:
        return f"❌ Search failed: {results['error']}"
    
    # Log usage
    _log_serpapi_search(f"Scholar: {query}", len(results))
    
    if not results:
        return "No results found."
    
    # Format output
    output = [f"🎓 Google Scholar Results for '{query}':\n"]
    
    for i, res in enumerate(results, 1):
        cited = f" (Cited by {res['cited_by']})" if res['cited_by'] else ""
        output.append(
            f"{i}. {res['title']}\n"
            f"   {res['publication_info']}{cited}\n"
            f"   Snippet: {res['snippet']}\n"
            f"   Link: {res['link']}"
        )
    
    return "\n\n".join(output)


google_scholar_tool = FunctionTool(func=search_google_scholar)


# ---------------------------------------------------------------------
# Cover Letter Generator - AI-powered with PDF & Word export
# ---------------------------------------------------------------------

def _create_word_document(cover_letter_text: str, candidate_name: str, output_path: str):
    """
    Create a professionally formatted Word document for the cover letter.
    
    Args:
        cover_letter_text (str): The cover letter content
        candidate_name (str): Candidate's name for header
        output_path (str): Path to save the .docx file
    
    Returns:
        str: Path to created document
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header with name and date
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"{candidate_name}\n")
    run.bold = True
    run.font.size = Pt(12)
    run = header.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(11)
    
    # Add spacing
    doc.add_paragraph()
    
    # Add cover letter body
    paragraphs = cover_letter_text.strip().split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_format = para.paragraph_format
            para_format.space_after = Pt(10)
            para_format.line_spacing = 1.15
            
            # Format paragraph text
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    doc.save(output_path)
    return output_path


def _convert_to_pdf(docx_path: str):
    """
    Convert a Word document to PDF.
    
    Args:
        docx_path (str): Path to the .docx file
    
    Returns:
        str: Path to created PDF file, or error message
    """
    try:
        from docx2pdf import convert
        pdf_path = docx_path.replace('.docx', '.pdf')
        convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e:
        # Fallback error message
        return f"PDF conversion failed: {str(e)}. Word document created successfully at {docx_path}"


async def generate_cover_letter(
    company_name: str,
    position_title: str,
    job_description: str,
    output_format: str = "both",
    cv_filename: str = "Professional Curriculum Vitae.docx",
    custom_notes: str = "",
    candidate_name: str = "Noah Haag"
):
    """
    Generate a personalized cover letter using LLM based on CV and job description.
    
    Args:
        company_name (str): Target company name
        position_title (str): Job title/position
        job_description (str): Full job description text
        output_format (str): Output format - "docx", "pdf", or "both" (default: "both")
        cv_filename (str): CV filename in documents folder (default: "Professional Curriculum Vitae.docx")
        custom_notes (str, optional): Additional points to emphasize in the cover letter
        candidate_name (str): Candidate's name for document header (default: "Noah Haag")
    
    Returns:
        str: Success message with file paths
    
    Example:
        >>> await generate_cover_letter(
                company_name="Google",
                position_title="ML Engineer",
                job_description="We are seeking...",
                output_format="both"
            )
        "Cover letter generated successfully!
        Word: cover_letters/Google_ML_Engineer_2025-11-24.docx
- DO NOT include placeholder text like [Your Name], [Date], [Company Address] - just write the body paragraphs
- Start directly with the salutation "Dear Hiring Manager," or similar

{f"CUSTOM NOTES TO EMPHASIZE: {custom_notes}" if custom_notes else ""}

Generate the cover letter now (body paragraphs only):"""

    # Import LLM from agent module to generate cover letter
    try:
        # Use litellm since it's already imported in the agent
        from google.adk.models.lite_llm import LiteLlm
        llm = LiteLlm(model="ollama_chat/llama3.2")
        
        response = await llm.generate_content(prompt)
        cover_letter_text = response.text.strip()
        
    except Exception as e:
        return f"Error generating cover letter with LLM: {str(e)}"
    
    # Generate filename
    safe_company = company_name.replace(" ", "_").replace("/", "-")
    safe_position = position_title.replace(" ", "_").replace("/", "-")
    date_str = datetime.now().strftime("%Y-%m-%d")
    base_filename = f"{safe_company}_{safe_position}_{date_str}"
    
    docx_path = os.path.join(COVER_LETTERS_FOLDER, f"{base_filename}.docx")
    
    # Create Word document
    try:
        _create_word_document(cover_letter_text, candidate_name, docx_path)
    except Exception as e:
        return f"Error creating Word document: {str(e)}"
    
    result_message = f"✅ Cover letter generated successfully!\n\nWord Document: {docx_path}"
    
    # Convert to PDF if requested
    if output_format in ["pdf", "both"]:
        pdf_result = await asyncio.to_thread(_convert_to_pdf, docx_path)
        
        if pdf_result.endswith(".pdf"):
            result_message += f"\nPDF Document: {pdf_result}"
        else:
            result_message += f"\n⚠️ {pdf_result}"
    
    # Update job application tracker if this application exists
    data = _load_job_applications()
    matching_apps = [app for app in data["applications"] 
                    if app["company"].lower() == company_name.lower() 
                    and app["position"].lower() == position_title.lower()]
    
    if matching_apps:
        # Update most recent matching application
        app_to_update = max(matching_apps, key=lambda x: x["last_updated"])
        app_to_update["cover_letter_generated"] = True
        app_to_update["last_updated"] = datetime.now().isoformat()
        _save_job_applications(data)
        result_message += f"\n\n📌 Updated application tracker for {company_name} - {position_title}"
    
    return result_message


# Wrap as ADK FunctionTool
cover_letter_generator_tool = FunctionTool(func=generate_cover_letter)


async def interview_prep_master(company_name: str, role_title: str = "Candidate"):
    """
    Generate a comprehensive interview preparation guide using LLM and CV context.
    
    Args:
        company_name (str): Name of the company.
        role_title (str): Title of the role (optional).
        
    Returns:
        str: Success message with path to the generated prep doc.
    """
    # Create prep directory
    prep_dir = os.path.join(os.getcwd(), "interview_prep")
    if not os.path.exists(prep_dir):
        os.makedirs(prep_dir)
        
    # Read CV and Resume
    cv_text = read_document("Professional Curriculum Vitae.docx")
    resume_text = read_document("Resume.docx")
    
    context = ""
    if not cv_text.startswith("Error"):
        context += f"\nCV Content:\n{cv_text}\n"
    if not resume_text.startswith("Error"):
        context += f"\nResume Content:\n{resume_text}\n"
        
    if not context:
        return "Error: Could not read CV or Resume for context."

    prompt = f"""
    You are an expert Interview Coach. Prepare a strategic interview preparation guide for {company_name}.
    
    CANDIDATE CONTEXT:
    {context}
    
    TARGET ROLE: {role_title}
    
    TASK:
    Generate a "Strategic Interview Prep Doc" in Markdown format.
    
    STRUCTURE:
    1. 🏢 **Company Strategy & Culture** (Infer from company name - mission, values, recent focus)
    2. 🎯 **The "Why Us?" Answer** (Connect candidate's specific background to the company)
    3. ❓ **Likely Technical Questions** (Based on candidate's skills vs company domain)
    4. 🗣️ **Behavioral Stories** (STAR method prompts based on candidate's actual history)
    5. 🧠 **3 "Killer" Questions to Ask Them** (Strategic questions that show deep insight)
    
    Keep it actionable, concise, and empowering.
    """
    
    try:
        from google.adk.models.lite_llm import LiteLlm
        llm = LiteLlm(model="ollama_chat/llama3.2")
        
        print(f"Generating interview prep for {company_name}...")
        response = await llm.generate_content(prompt)
        content = response.text.strip()
        
        # Save to file
        safe_name = company_name.replace(" ", "_").replace("/", "-")
        filename = f"Prep_{safe_name}_{datetime.now().strftime('%Y-%m-%d')}.md"
        filepath = os.path.join(prep_dir, filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
            
        return f"✅ Interview Prep Guide generated: {filepath}"
        
    except Exception as e:
        return f"Error generating prep guide: {str(e)}"


interview_prep_tool = FunctionTool(func=interview_prep_master)


def read_document(filename: str = "Professional Curriculum Vitae.docx"):
    """
    Reads a document (PDF or Word) from the local 'documents' folder and returns its text content.

    If no filename is provided, it defaults to reading
    'Professional Curriculum Vitae.docx'.

    Supported file types: .pdf, .docx

    Args:
        filename (str, optional): The name of the document to read. Defaults to None.

    Returns:
        str: Extracted text from the document, or an error message if the file doesn't exist
             or is an unsupported type.
    """
    if filename is None:
        filename = "Professional Curriculum Vitae.docx"

    filepath = os.path.join(DOCUMENT_FOLDER, filename)

    if not os.path.exists(filepath):
        return f"Error: File '{filename}' does not exist in the documents folder."

    try:
        if filename.lower().endswith(".pdf"):
            reader = PdfReader(filepath)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif filename.lower().endswith(".docx"):
            doc = docx.Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            return "Error: Unsupported file type. Only .pdf and .docx are allowed."

        return text.strip() if text else "No text found in document."
    except Exception as e:
        return f"Failed to read document: {e}"


papers_path = "chroma_db_research_papers_test"

try:
    embedding_func = get_embedding_function()
    if embedding_func is None:
        print("[WARNING] Skipping ChromaDB initialization: No embedding function (GitHub Actions?)")
        db_papers = None
    elif not os.path.exists(papers_path):
        print(f"[WARNING] Skipping ChromaDB initialization: Path not found {papers_path}")
        db_papers = None
    else:
        db_papers = Chroma(persist_directory=papers_path, embedding_function=embedding_func,
                           client_settings=Settings(anonymized_telemetry=False))
except Exception as e:
    print(f"[WARNING] Failed to initialize ChromaDB: {e}")
    db_papers = None


def search_pdf(query: str):
    if db_papers is None:
        return "Error: Research paper database is not available in this environment."
        
    retriever = db_papers.as_retriever(search_kwargs={"k": 25,
                                                      "fetch_k": 250,
                                                      "lambda_mult": 0.8},
                                       search_type="mmr")

    return retriever.invoke(query)


def create_gmail_tools(
        token_file=r"D:\Python Projects\AI Agents\Agent_V2\token.json",
        client_secrets_file=r"D:\Python Projects\AI Agents\Agent_V2\credentials.json",
        scopes=None
):
    """
    Initializes Gmail credentials, builds the Gmail service,
    and returns the Gmail toolkit tools.
    """

    if scopes is None:
        scopes = ["https://mail.google.com/"]

    # Step 1: Authenticate
    credentials = get_google_credentials(
        token_file=token_file,
        scopes=scopes,
        client_secrets_file=client_secrets_file,
    )

    # Step 2: Build Gmail API service
    api_resource = build_gmail_service(credentials=credentials)

    # Step 3: Create toolkit
    toolkit = GmailToolkit(api_resource=api_resource)

    # Step 4: Return the tool(s)
    return toolkit.get_tools()


# ---------------------------------------------------------------------
# Load Gmail tools (optional - gracefully handle missing credentials)
# ---------------------------------------------------------------------
try:
    gmail_tools = create_gmail_tools(
        token_file="token.json",
        client_secrets_file="credentials.json"
    )

    # Filter out any sending tools
    gmail_tools_filtered = [
        t for t in gmail_tools if "send" not in t.name.lower()
    ]

    gmail_draft_tool = next(t for t in gmail_tools if t.name == "create_gmail_draft")
    gmail_get_msg_tool = next(t for t in gmail_tools if t.name == "get_gmail_message")
    gmail_search_tool = next(t for t in gmail_tools if t.name == "search_gmail")
    
    print("[INFO] Gmail tools initialized successfully")
    GMAIL_AVAILABLE = True
    
except (FileNotFoundError, Exception) as e:
    print(f"[WARNING] Gmail tools not available: {e}")
    print("[INFO] Agent will run without Gmail functionality")
    
    # Create dummy tools that return helpful error messages
    gmail_draft_tool = None
    gmail_get_msg_tool = None
    gmail_search_tool = None
    gmail_tools_filtered = []
    GMAIL_AVAILABLE = False


# print("[DEBUG] Draft tool selected:", gmail_draft_tool.name)
# print("[DEBUG] Search tool selected:", gmail_search_tool.name)
# print("[DEBUG] GetMessage tool selected:", gmail_get_msg_tool.name)

# ---------------------------------------------------------------------
# Async wrapper around the Gmail draft tool
# ---------------------------------------------------------------------
async def draft_email_async(to: str, subject: str, body: str):
    """
    Asynchronously creates a Gmail draft using the Gmail draft tool returned
    by the `create_gmail_tools()` GmailToolkit loader.

    This function is an async wrapper around the synchronous
    `gmail_draft_tool.run()` call. Because GmailToolkit tools execute using
    the Google API client (which is blocking), this coroutine uses
    `asyncio.to_thread()` to keep the agent's event loop responsive while
    still leveraging the synchronous Gmail API bindings.

    Relationship to the Gmail Toolkit
    ---------------------------------
    - `create_gmail_tools()` authenticates the user, builds the Gmail API
      resource, constructs a GmailToolkit, and returns a list of Tool objects.
    - `gmail_draft_tool` is selected from this list via:
            gmail_draft_tool = next(t for t in gmail_tools if t.name == "create_gmail_draft")
    - This async wrapper simply exposes that tool in a form that agents can
      safely call within an asynchronous workflow.

    Parameters
    ----------
    to : str
        The destination email address for the draft. This function wraps
        the address in a one-element list because the GmailToolkit's
        draft tool expects:  {"to": ["recipient@domain.com"], ...}.

    subject : str
        The email subject line.

    body : str
        The body of the message. This is treated as plain-text unless the
        underlying GmailToolkit implementation supports HTML bodies.

    Returns
    -------
    dict
        The complete Gmail API response returned by
        `gmail_draft_tool.run(tool_input)`. This usually includes metadata
        such as:
            - draft ID
            - message ID
            - internal Gmail payload content

    Notes
    -----
    - This function *does not send email*. It only creates a draft.
    - All Gmail operations in this project come from the shared GmailToolkit
      instance produced by `create_gmail_tools()`.
    - Because the toolkit uses synchronous Google API calls, using
      `asyncio.to_thread()` prevents blocking your LLM agent runtime.
    - Debug prints are included to trace internal behavior.

    Example
    -------
     await draft_email_async(
            to="someone@example.com",
            subject="Follow-Up",
            body="Thanks again for meeting today!"
        )
    {'id': 'draft_12345', 'message': {...}}
    """
    # Check if Gmail is available
    if not GMAIL_AVAILABLE or gmail_draft_tool is None:
        return "Error: Gmail functionality is not available. This requires Gmail credentials (credentials.json and token.json) which are only available when running locally."

    print(f"[DEBUG] Starting draft_email_async: to={to}, subject={subject}")

    tool_input = {
        "to": [to],
        "subject": subject,
        "message": body,
    }

    # Run GmailCreateDraft inside thread
    result = await asyncio.to_thread(gmail_draft_tool.run, tool_input)

    print(f"[DEBUG] Finished draft_email_async: result={result}")
    return result


gmail_draft_tool_for_agent = FunctionTool(
    func=draft_email_async
)


# ---------------------------------------------------------------------
# Utility to clean extracted email text
# ---------------------------------------------------------------------
def clean_email_text(raw_text: str) -> str:
    """
    Normalise whitespace, strip URLs and trim the result.
    """
    # Collapse multiple whitespace / newlines
    text = re.sub(r"\s+", " ", raw_text)

    # Remove URLs (tracking, “read more”, unsubscribe links, etc.)
    text = re.sub(r"https?://\S+", "", text)

    # Strip leading / trailing whitespace
    return text.strip()


# ---------------------------------------------------------------------
# Helper: safely extract a readable body from a Gmail message part
# ---------------------------------------------------------------------
def _extract_message_body(part: dict) -> str:
    """
    Return plain‑text content from a Gmail API message part.

    - Skips multipart containers and attachments.
    - Prefers ``text/plain``; falls back to ``text/html`` (converted with BeautifulSoup).
    - Tries UTF‑8 → latin‑1 decoding, finally decodes with ``errors="ignore"``.
    """
    # Containers / attachments have no useful body
    mime_type = part.get("mimeType", "")
    if mime_type.startswith("multipart/"):
        return ""

    if part.get("filename"):          # attachment
        return ""

    # Grab the raw base64 data (URL‑safe) if present
    body_data = part.get("body", {}).get("data")
    if not body_data:
        return ""

    # Gmail sometimes omits padding; add it back
    padding = "=" * (-len(body_data) % 4)
    try:
        decoded = base64.urlsafe_b64decode(body_data + padding).decode("utf-8")
    except UnicodeDecodeError:
        # Try latin‑1, then fallback to ignore‑errors
        try:
            decoded = base64.urlsafe_b64decode(body_data + padding).decode("latin-1")
        except Exception:
            decoded = base64.urlsafe_b64decode(body_data + padding).decode(
                "utf-8", errors="ignore"
            )

    # Plain‑text: just strip
    if mime_type == "text/plain":
        return decoded.strip()

    # HTML: convert to plain text
    if mime_type == "text/html":
        from bs4 import BeautifulSoup

        return BeautifulSoup(decoded, "html.parser").get_text().strip()

    return ""


# ---------------------------------------------------------------------
# Synchronous email reader for the agent
# ---------------------------------------------------------------------
def read_emails_for_agent(query: str = "", max_results: int = 5) -> dict:
    """
    Search Gmail, fetch each message, and return a clean plain‑text version.

    Returns
    -------
    dict
        {"messages": [{"id": "...", "snippet": "...", "text": "..."}]}
    """
    # Check if Gmail is available
    if not GMAIL_AVAILABLE or gmail_search_tool is None:
        return {"error": "Gmail functionality is not available. This requires Gmail credentials (credentials.json and token.json) which are only available when running locally."}
    
    print(f"[DEBUG] read_emails_for_agent called: query='{query}', max={max_results}")

    # 1️⃣  Perform the Gmail search
    search_input = {"query": query, "maxResults": max_results}
    search_results = gmail_search_tool.run(search_input)

    # Normalise the result to a list of message dicts
    if isinstance(search_results, list):
        messages_list = search_results
    elif isinstance(search_results, dict):
        messages_list = search_results.get("messages", [])
    else:
        print("[DEBUG] Unexpected search_results type")
        return {"messages": []}

    print(f"[DEBUG] messages_list length: {len(messages_list)}")
    if not messages_list:
        return {"messages": []}

    messages = []

    # 2️⃣  Iterate over each message ID and fetch full payload
    for msg in messages_list:
        msg_id = msg.get("id")
        if not msg_id:
            continue

        msg_detail = gmail_get_msg_tool.run({"message_id": msg_id})
        snippet = msg_detail.get("snippet", "")
        text_parts = [snippet] if snippet else []

        # 3️⃣  Recursively walk the payload tree and extract text
        def walk(part: dict):
            texts = [_extract_message_body(part)]
            for sub in part.get("parts", []):
                texts.append(walk(sub))
            # flatten any nested lists
            return "\n".join(filter(None, texts))

        payload = msg_detail.get("payload", {})
        extracted = walk(payload)
        if extracted:
            text_parts.append(extracted)

        full_text = "\n".join(filter(None, text_parts))
        clean_text = clean_email_text(full_text)

        messages.append(
            {"id": msg_id, "snippet": snippet, "text": clean_text or "<NO TEXT FOUND>"}
        )

    print(f"[DEBUG] Total messages fetched: {len(messages)}")
    return {"messages": messages}


# ---------------------------------------------------------------------
# Wrap as FunctionTool for the agent (synchronous)
# ---------------------------------------------------------------------
gmail_read_tool_for_agent = FunctionTool(
    func=read_emails_for_agent,
    require_confirmation=False,
)
